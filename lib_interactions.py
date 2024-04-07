import nltk
import docx
import re
import string
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.tokenize import sent_tokenize


UNSHIPHER = {
    "CC": "союз",
    "CD": "кардинальное число",
    "DT": "определитель",
    "EX": 'существование там',
    "FW": "иностранное слово",
    "IN": "предлог/подчинительный союз",
    "JJ": 'прилагательное',
    "VP": "глагольная группа",
    "JJR": 'прилагательное, сравнительная степень',
    "JJS": 'прилагательное, превосходная степень',
    "LS": "маркер списка  1)",
    "MD": "модальный глагол сосотавное сказуемое",
    "NN": 'существительное, единственное число',
    "NNS": 'существительное, множественное число',
    "PP": "предложная группа",
    "NNP": 'имя собственное, единственное число',
    "NNPS": 'имя собственное, множественное число',
    "PDT": 'предопределитель',
    "POS": "притяжательное окончание",
    "PRP": "личное местоимение, ",
    "PRP$": "притяжательное местоимение",
    "RB": "наречие",
    "RBR": "наречие, сравнительная степень",
    "RBS": "наречие, превосходная степень",
    "RP": "частица",
    "SBAR": "Предложение, введенное (возможно пустым) подчинительным союзом",
    "SBARQ": "Прямой вопрос, введенный вопросительным словом или вопросительной группой",
    "SINV": "Инвертированное повествовательное предложение, т.е. такое, в котором подлежащее следует за глаголом в прошедшем времени или модальным глаголом.",
    "SQ": "Инвертированный вопрос да/нет, или главное предложение вопроса, следующее за вопросительной группой в SBARQ",
    "SYM": "Символ",
    "VBD": "глагол, прошедшее время",
    "VBG": "глагол, герундий/презенс-партицип  берущий",
    "VBN": "глагол, прошедшее причастие  взятый",
    "VBP": "глагол, настоящее время, ед. число, не 3-е лицо",
    "VBZ": "глагол, настоящее время, 3-е лицо, ед. число",
    "WDT": "вопросительный определитель",
    "WP": "вопросительное местоимение",
    "WP$": "притяжательное вопросительное местоимение",
    "WRB": "вопросительное наречие",
    "TO": 'to',
    "UH": "междометие",
    "VB": "глагол, исходная форма",
}
signs = "!~@#$%^&*()_+<>?:.,’;[]\\|'\"\'–«‘1234567890'”`“"


def load_file(path: str) -> list:
    opened_file = open(path, "rb")
    if '.doc' in path:
        current_file = docx.Document(opened_file)
        list_of_rows = [row.text for row in current_file.paragraphs]
        opened_file.close()
        raw_data = ""
        for row in list_of_rows:
            raw_data += row
    elif '.txt' in path:
        raw_data = ''
        for raw in opened_file:
            raw_data += raw
    else:
        raw_data = None
    return raw_data


def process_text(raw_data):

    sentences = sent_tokenize(raw_data)
    for sindex in range(len(sentences)):
        for sign in signs:
            if sign in sentences[sindex]:
                sentences[sindex] = sentences[sindex].replace(sign, ' ')
    sentence_dict = {sentence: [] for sentence in sentences}
    grammar = nltk.RegexpParser('''
    NP: {<DT>?<JJ>*<NN.*>}
    P: {<IN>}           
    V: {<V.*>}          
    PP: {<P> <NP>}      
    VP: {<V> <NP|PP>*}  
    ADJP: {<JJ>}        
    S: {<NP> <VP>}      
    ''')
    for sentence in sentences:
        tree = grammar.parse(nltk.pos_tag(
            word_tokenize(sentence)))
        sentence_dict[sentence].append(tree)
    result = []
    for sentence in sentence_dict.keys():
        result.append((sentence, sentence_dict.get(sentence)))
    return result


def save_file(text: list, name: str):
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "id"
    hdr_cells[1].text = "word"
    hdr_cells[2].text = "description"
    for id, word, description in text:
        if not description:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = str(id)
        row_cells[1].text = word
        row_cells[2].text = description

    document.add_page_break()

    document.save(name + ".docx")
